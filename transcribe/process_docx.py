#!/usr/bin/env python3
"""Process DOCX templates: extract structure, fill with content, read references."""

import argparse
import copy
import json
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _iter_body(doc):
    """Yield (type_str, object) for each body element in document order."""
    tbl_map = {t._element: t for t in doc.tables}
    para_map = {p._element: p for p in doc.paragraphs}
    for child in doc.element.body:
        if child in tbl_map:
            yield "table", tbl_map[child]
        elif child in para_map:
            yield "paragraph", para_map[child]


def _fmt_runs(para):
    """Collect runs with formatting info."""
    runs = []
    for r in para.runs:
        if not r.text:
            continue
        fmt = []
        if r.bold:
            fmt.append("bold")
        if r.italic:
            fmt.append("italic")
        runs.append((r.text, fmt))
    return runs


# ---------------------------------------------------------------------------
# --extract
# ---------------------------------------------------------------------------

def extract(template_path):
    """Extract and print template structure in human-readable form."""
    doc = Document(template_path)
    print(f"TEMPLATE: {Path(template_path).name}")
    print("=" * 30)
    print()

    tbl_idx = 0
    for kind, obj in _iter_body(doc):
        if kind == "table":
            tbl = obj
            n_rows = len(tbl.rows)
            n_cols = len(tbl.columns)
            print(f"[Table {tbl_idx}] {n_cols} cols x {n_rows} rows:")

            for ri, row in enumerate(tbl.rows):
                cells = []
                for ci in range(len(row.cells)):
                    cells.append(row.cells[ci].text.replace("\n", "\\n").strip())

                if n_cols >= 4:
                    print(f"  ({ri}): {' | '.join(cells)}")
                else:
                    parts = [f'({ri},{ci}): "{cells[ci]}"' for ci in range(len(cells))]
                    print(f"  {'  '.join(parts)}")

            tbl_idx += 1
            print()

        elif kind == "paragraph":
            para = obj
            text = para.text
            if not text.strip():
                continue

            style = para.style.name if para.style else "Normal"
            runs = _fmt_runs(para)

            if style.startswith("Heading"):
                level = style.replace("Heading ", "")
                print(f'[Heading {level}] "{text}"')

            elif style == "List Paragraph":
                parts = []
                for rt, rf in runs:
                    if rf:
                        parts.append(f'{", ".join(rf)}:"{rt}"')
                    else:
                        parts.append(f'"{rt}"')
                print(f'[ListParagraph] {" ".join(parts)}')

            else:
                fmts = set(tuple(rf) for _, rf in runs)
                if len(fmts) > 1:
                    parts = []
                    for rt, rf in runs:
                        if rf:
                            parts.append(f'{", ".join(rf)}: "{rt}"')
                        else:
                            parts.append(f'"{rt}"')
                    print(f'[Para] style={style}, {", ".join(parts)}')
                elif len(fmts) == 1 and fmts != {()}:
                    fmt_str = ", ".join(runs[0][1])
                    print(f'[Para] style={style}, {fmt_str}: "{text}"')
                else:
                    print(f'[Para] style={style}: "{text}"')


# ---------------------------------------------------------------------------
# --references
# ---------------------------------------------------------------------------

def references(dir_path):
    """Read all .docx files in directory and output their text content."""
    d = Path(dir_path)
    if not d.is_dir():
        print(f"Error: {dir_path} is not a directory", file=sys.stderr)
        sys.exit(1)

    docx_files = sorted(d.glob("*.docx"))
    if not docx_files:
        print(f"No .docx files found in {dir_path}")
        return

    for f in docx_files:
        print(f"REFERENCE: {f.name}")
        print("=" * 30)
        doc = Document(str(f))

        for kind, obj in _iter_body(doc):
            if kind == "paragraph":
                text = obj.text.strip()
                if text:
                    print(text)
            elif kind == "table":
                for row in obj.rows:
                    cells = [row.cells[ci].text.strip() for ci in range(len(row.cells))]
                    print(" | ".join(cells))
                print()

        print()
        print()


# ---------------------------------------------------------------------------
# --fill
# ---------------------------------------------------------------------------

def fill(template_path, content_path, output_path):
    """Fill template with content from JSON, preserving formatting."""
    with open(content_path, "r", encoding="utf-8") as f:
        content = json.load(f)

    # Save template table formatting
    tpl_doc = Document(template_path)
    tpl_tables = [copy.deepcopy(t._element) for t in tpl_doc.tables]

    # Save paragraph properties from template cells
    tpl_cell_pPr = {}
    for ti, tbl in enumerate(tpl_doc.tables):
        tpl_cell_pPr[ti] = {}
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                if cell.paragraphs:
                    pPr = cell.paragraphs[0]._element.find(qn("w:pPr"))
                    if pPr is not None:
                        tpl_cell_pPr[ti][(ri, ci)] = copy.deepcopy(pPr)

    # Copy template file (preserves page setup, headers, footers, styles)
    shutil.copy2(template_path, output_path)
    doc = Document(output_path)

    # Clear body (keep sectPr for page layout)
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split("}")[-1]
        if tag in ("p", "tbl"):
            body.remove(child)

    # Build document from content
    for el in content.get("elements", []):
        el_type = el.get("type")
        if el_type == "paragraph":
            _add_para(doc, el)
        elif el_type == "table":
            _add_table(doc, el, tpl_tables, tpl_cell_pPr)

    doc.save(output_path)
    print(f"Created: {output_path}")


def _resolve_style(doc, style_name):
    """Find a style by name or style_id, return None if not found."""
    from docx.enum.style import WD_STYLE_TYPE

    # Try by name
    try:
        return doc.styles[style_name]
    except KeyError:
        pass
    # Try by style_id (e.g. "Heading1" for "Heading 1")
    try:
        return doc.styles.get_by_id(
            style_name.replace(" ", ""), WD_STYLE_TYPE.PARAGRAPH
        )
    except Exception:
        pass
    return None


def _add_para(doc, el):
    """Add a paragraph element to the document."""
    style_name = el.get("style", "Normal")
    para = doc.add_paragraph()
    style = _resolve_style(doc, style_name)
    if style is not None:
        para.style = style

    if "runs" in el:
        for rd in el["runs"]:
            run = para.add_run(rd.get("text", ""))
            if rd.get("bold"):
                run.bold = True
            if rd.get("italic"):
                run.italic = True
    elif "text" in el:
        para.add_run(el["text"])


def _copy_xml_prop(parent, tpl_parent, tag_name, insert_pos=0):
    """Copy an XML property element from template to target, replacing if exists."""
    tpl_el = tpl_parent.find(qn(tag_name))
    if tpl_el is None:
        return
    new_el = copy.deepcopy(tpl_el)
    old_el = parent.find(qn(tag_name))
    if old_el is not None:
        parent.replace(old_el, new_el)
    else:
        parent.insert(insert_pos, new_el)


def _add_table(doc, el, tpl_tables, tpl_cell_pPr):
    """Add a table element with optional template formatting."""
    rows_data = el.get("rows", [])
    if not rows_data:
        return

    n_rows = len(rows_data)
    n_cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=n_rows, cols=n_cols)

    tpl_idx = el.get("template_table")
    if tpl_idx is not None and tpl_idx < len(tpl_tables):
        tpl = tpl_tables[tpl_idx]

        # Copy table properties (borders, style, widths)
        _copy_xml_prop(table._element, tpl, "w:tblPr", 0)

        # Copy column grid
        tpl_grid = tpl.find(qn("w:tblGrid"))
        if tpl_grid is not None:
            new_grid = copy.deepcopy(tpl_grid)
            old_grid = table._element.find(qn("w:tblGrid"))
            if old_grid is not None:
                table._element.replace(old_grid, new_grid)

        # Copy row and cell properties
        tpl_rows = tpl.findall(qn("w:tr"))
        for ri, row in enumerate(table.rows):
            tri = min(ri, len(tpl_rows) - 1) if tpl_rows else None
            if tri is None:
                continue

            tpl_row = tpl_rows[tri]

            # Row properties
            _copy_xml_prop(row._tr, tpl_row, "w:trPr", 0)

            # Cell properties
            tpl_cells = tpl_row.findall(qn("w:tc"))
            for ci, cell in enumerate(row.cells):
                tci = min(ci, len(tpl_cells) - 1) if tpl_cells else None
                if tci is None:
                    continue

                _copy_xml_prop(cell._tc, tpl_cells[tci], "w:tcPr", 0)

                # Paragraph properties from template cell
                pPr_key = (tri, tci)
                if tpl_idx in tpl_cell_pPr and pPr_key in tpl_cell_pPr[tpl_idx]:
                    new_pPr = copy.deepcopy(tpl_cell_pPr[tpl_idx][pPr_key])
                    if cell.paragraphs:
                        p_el = cell.paragraphs[0]._element
                        old_pPr = p_el.find(qn("w:pPr"))
                        if old_pPr is not None:
                            p_el.replace(old_pPr, new_pPr)
                        else:
                            p_el.insert(0, new_pPr)

    # Fill cells with content
    for ri, row_data in enumerate(rows_data):
        for ci, cell_data in enumerate(row_data):
            if ci >= n_cols:
                continue
            cell = table.cell(ri, ci)

            if isinstance(cell_data, str):
                cell_data = {"text": cell_data}

            para = cell.paragraphs[0]

            if "runs" in cell_data:
                for rd in cell_data["runs"]:
                    text = rd.get("text", "")
                    lines = text.split("\n")
                    for i, line in enumerate(lines):
                        run = para.add_run(line)
                        if rd.get("bold"):
                            run.bold = True
                        if rd.get("italic"):
                            run.italic = True
                        if i < len(lines) - 1:
                            run.add_break()
            else:
                text = cell_data.get("text", "")
                lines = text.split("\n")
                for i, line in enumerate(lines):
                    run = para.add_run(line)
                    if i < len(lines) - 1:
                        run.add_break()


# ---------------------------------------------------------------------------
# --transcribe
# ---------------------------------------------------------------------------

def transcribe(audio_path, output_dir=None):
    """Transcribe audio file using mlx-whisper (Apple Silicon GPU), save as .txt."""
    import mlx_whisper

    audio = Path(audio_path)
    if not audio.exists():
        print(f"Error: {audio_path} not found", file=sys.stderr)
        sys.exit(1)

    out_dir = Path(output_dir) if output_dir else audio.parent
    out_file = out_dir / f"{audio.stem}.txt"

    print(f"Transcribing: {audio.name} ...")
    result = mlx_whisper.transcribe(
        str(audio),
        path_or_hf_repo="mlx-community/whisper-large-v3-mlx",
        verbose=False,
    )

    lang = result.get("language", "unknown")
    print(f"Detected language: {lang}")

    with open(out_file, "w", encoding="utf-8") as f:
        for segment in result.get("segments", []):
            f.write(segment["text"].strip() + "\n")

    print(f"Transcript saved: {out_file}")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Process DOCX templates")
    parser.add_argument(
        "--extract", metavar="TEMPLATE",
        help="Extract and print template structure",
    )
    parser.add_argument(
        "--fill", nargs=3, metavar=("TEMPLATE", "CONTENT", "OUTPUT"),
        help="Fill template with content JSON",
    )
    parser.add_argument(
        "--references", metavar="DIR",
        help="Read all .docx reference files from directory",
    )
    parser.add_argument(
        "--transcribe", metavar="AUDIO",
        help="Transcribe audio file to .txt",
    )
    parser.add_argument(
        "--output_dir", metavar="DIR",
        help="Output directory for transcription (default: same as audio)",
    )
    args = parser.parse_args()

    if args.extract:
        extract(args.extract)
    elif args.fill:
        fill(*args.fill)
    elif args.references:
        references(args.references)
    elif args.transcribe:
        transcribe(args.transcribe, args.output_dir)
    else:
        parser.print_help()
        sys.exit(1)


if __name__ == "__main__":
    main()
